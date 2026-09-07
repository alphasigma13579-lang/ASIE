from __future__ import annotations

from copy import deepcopy
import tempfile
import unittest
from pathlib import Path

from backend import asie_local_api as api
from backend.funder_report import build_funder_report_projection
from backend.funder_report import render_funder_report_html
from backend.customer_presentation import business_text, safe_narrative, text
from backend.customer_report_content import customer_report_groups
from backend.funding_readiness import evaluate_funding_readiness, profile_ids, sector_profile_catalog
from backend.report_release import build_release_record, validate_release_record
from backend.report_exports import export_funder_report_docx, export_funder_report_pdf, export_funder_report_pptx
from backend.snapshot_assembly import canonical_hash


class FunderReportProjectionTests(unittest.TestCase):
    def make_repo(self) -> api.Repository:
        temp_dir = tempfile.TemporaryDirectory()
        self.addCleanup(temp_dir.cleanup)
        return api.Repository(Path(temp_dir.name) / "asie-funder-test.sqlite3")

    def test_projection_is_snapshot_bound_and_has_sixteen_sections(self) -> None:
        repo = self.make_repo()
        project = repo.create_project(
            {
                "name": "حزمة تمويل تجريبية",
                "sector": "خدمات",
                "jurisdiction": "Saudi Arabia",
                "inputs": {
                    "startup_cost": 250000,
                    "monthly_fixed_cost": 62000,
                    "unit_price": 85,
                    "variable_cost": 34,
                    "monthly_units": 1600,
                    "annual_discount_rate": 0.10,
                    "working_capital_months": 2,
                    "equity_contribution": 150000,
                    "debt_amount": 0,
                    "loan_grace_months": 0,
                },
            }
        )
        overview, report = api.build_overview(project, repo)
        projection = report["funder_report"]

        self.assertEqual(projection["contract_id"], "funder.report.projection.v1")
        self.assertEqual(projection["snapshot_id"], overview["snapshot"]["snapshot_id"])
        self.assertEqual(len(projection["sections"]), 16)
        self.assertEqual(projection["sections"][13]["section_id"], "14-financial-expectations")
        self.assertEqual(projection["sections"][14]["section_id"], "15-capital-requirements")
        self.assertEqual(projection["sections"][13]["payload"]["statements"]["status"], "partial")
        self.assertIn("balance_sheet", projection["gaps"])

        unhashed = dict(projection)
        unhashed.pop("projection_hash")
        self.assertEqual(projection["projection_hash"], canonical_hash(unhashed))

    def test_projection_exposes_input_traceability_bound_to_same_snapshot(self) -> None:
        repo = self.make_repo()
        project = repo.create_project({"name": "تتبع المدخلات", "inputs": {"startup_cost": 250000}})
        overview, report = api.build_overview(project, repo)
        traceability = report["funder_report"]["input_traceability"]
        self.assertEqual("input.traceability.v1", traceability["contract_id"])
        self.assertEqual(overview["snapshot"]["snapshot_id"], traceability["snapshot_id"])
        startup = next(row for row in traceability["items"] if row["input_key"] == "startup_cost")
        self.assertEqual("user_input", startup["source_type"])
        self.assertEqual("draft", startup["review_status"])

    def test_projection_does_not_recalculate_finance(self) -> None:
        repo = self.make_repo()
        project = repo.create_project({"name": "مدخلات ناقصة", "inputs": {}})
        overview, report = api.build_overview(project, repo)
        projection = build_funder_report_projection(overview)

        self.assertEqual(projection["readiness_status"], "DRAFT_INTERNAL")
        self.assertEqual(projection["sections"][13]["payload"]["statements"]["status"], "not_ready")
        self.assertEqual(report["funder_report"]["snapshot_id"], overview["snapshot"]["snapshot_id"])

    def test_demo_projection_is_explicitly_blocked_from_production(self) -> None:
        repo = self.make_repo()
        project = repo.create_project({"name": "بيانات تجريبية", "inputs": {}})
        overview, report = api.build_overview(project, repo)
        projection = report["funder_report"]
        self.assertEqual("demo_simulated_external", projection["data_mode"])
        self.assertEqual("DEMO / LOCAL ONLY", projection["display_badge"])
        self.assertEqual("blocked", projection["production_admission"])
        self.assertEqual("DRAFT_INTERNAL", projection["readiness_status"])
        self.assertIn("demo_data_not_admitted_to_production", projection["gaps"])

    def test_html_composer_is_localized_and_hides_internal_identifiers(self) -> None:
        repo = self.make_repo()
        project = repo.create_project({"name": "عرض تمويلي", "inputs": {}})
        overview, report = api.build_overview(project, repo)
        arabic = render_funder_report_html(report["funder_report"], locale="ar")
        english = render_funder_report_html(report["funder_report"], locale="en")

        self.assertIn("lang='ar' dir='rtl'", arabic)
        self.assertIn("تقرير جدوى المشروع", arabic)
        self.assertIn("ما الذي يحتاج استكمالاً؟", arabic)
        self.assertIn("lang='en' dir='ltr'", english)
        self.assertIn("Project feasibility report", english)
        self.assertIn("What still needs completion?", english)
        for output in (arabic, english):
            self.assertNotIn(overview["snapshot"]["snapshot_id"], output)
            self.assertNotIn(report["funder_report"]["contract_id"], output)
            self.assertNotIn(report["funder_report"]["projection_hash"], output)

    def test_customer_vocabulary_explains_known_codes_and_actions_in_both_languages(self) -> None:
        self.assertEqual("المصروفات التشغيلية مرتفعة مقارنة بالإيراد", business_text("opex_above_60_percent_of_revenue", "ar"))
        self.assertEqual("Operating costs are high relative to revenue", business_text("opex_above_60_percent_of_revenue", "en"))
        self.assertEqual(
            "خفّض المصروفات الثابتة أو أثبت قدرة أعلى على تحقيق الإيراد.",
            safe_narrative("Reduce fixed OPEX or increase validated revenue capacity.", "ar"),
        )
        self.assertEqual(
            "Earnings before interest, tax, depreciation, and amortisation",
            text("ebitda", "en"),
        )

    def test_docx_export_hides_snapshot_and_contains_customer_sections(self) -> None:
        try:
            from docx import Document
        except ModuleNotFoundError:
            self.skipTest("docx runtime is supplied by the document export environment")
        repo = self.make_repo()
        project = repo.create_project({"name": "حزمة Word", "inputs": {}})
        overview, report = api.build_overview(project, repo)
        with tempfile.TemporaryDirectory() as temp_dir:
            path = export_funder_report_docx(report["funder_report"], Path(temp_dir) / "funder.docx")
            self.assertTrue(path.exists())
            document = Document(path)
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            self.assertIn("تقرير جدوى المشروع", text)
            self.assertIn("ما الذي يحتاج استكمالاً؟", text)
            document_text = text + "\n" + "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
            self.assertNotIn(overview["snapshot"]["snapshot_id"], document_text)
            self.assertNotIn(report["funder_report"]["contract_id"], document_text)
            import zipfile
            with zipfile.ZipFile(path) as archive:
                self.assertIn("<w:bidi", archive.read("word/document.xml").decode("utf-8"))

            english_path = export_funder_report_docx(report["funder_report"], Path(temp_dir) / "funder-en.docx", locale="en")
            with zipfile.ZipFile(english_path) as archive:
                english_xml = archive.read("word/document.xml").decode("utf-8")
                self.assertNotIn("<w:bidi", english_xml)
                self.assertIn('w:jc w:val="left"', english_xml)

    def test_pdf_export_is_server_side_and_snapshot_bound(self) -> None:
        import os
        import shutil
        renderer = os.environ.get("ASIE_PDF_RENDERER") or shutil.which("chrome") or shutil.which("msedge")
        if renderer is None:
            self.skipTest("server-side PDF renderer is pinned in the production image")
        repo = self.make_repo()
        project = repo.create_project({"name": "حزمة PDF", "inputs": {}})
        overview, report = api.build_overview(project, repo)
        with tempfile.TemporaryDirectory() as temp_dir:
            path = export_funder_report_pdf(report["funder_report"], Path(temp_dir) / "funder.pdf")
            self.assertTrue(path.exists())
            self.assertGreater(path.stat().st_size, 1000)
            self.assertEqual(path.read_bytes()[:5], b"%PDF-")

    def test_pptx_export_is_openxml_and_snapshot_bound(self) -> None:
        import zipfile
        repo = self.make_repo()
        project = repo.create_project({"name": "حزمة PowerPoint", "inputs": {}})
        _overview, report = api.build_overview(project, repo)
        with tempfile.TemporaryDirectory() as temp_dir:
            path = export_funder_report_pptx(report["funder_report"], Path(temp_dir) / "funder.pptx", locale="en")
            with zipfile.ZipFile(path) as archive:
                self.assertIn("[Content_Types].xml", archive.namelist())
                self.assertIn("ppt/presentation.xml", archive.namelist())
                self.assertIn("ppt/slides/slide1.xml", archive.namelist())
                slides = "\n".join(archive.read(name).decode("utf-8") for name in archive.namelist() if name.startswith("ppt/slides/slide") and name.endswith(".xml"))
                self.assertIn("Project feasibility report", slides)
                self.assertIn('algn="l"', slides)
                self.assertNotIn('rtl="1"', slides)
                self.assertNotIn(report["funder_report"]["snapshot_id"], slides)
                self.assertNotIn(report["funder_report"]["contract_id"], slides)

    def test_document_formats_include_saved_business_content_without_mutation(self) -> None:
        """Content parity, not merely translated headings or valid file signatures."""
        import zipfile
        from xml.etree import ElementTree
        projection = {
            "snapshot_id": "PRIVATE-SNAPSHOT", "projection_hash": "PRIVATE-HASH",
            "readiness_status": "DRAFT_INTERNAL", "gaps": ["marketing_strategy"],
            "sections": [
                {"section_id": "02-executive-summary", "payload": {
                    "decision": {"sovereign_verdict": "revise_and_reassess", "reason": "negative_npv"},
                    "kpis": [{"output_id": "monthly_revenue", "value": 12345.6789, "unit": "SAR", "status": "ready"}],
                }},
                {"section_id": "12-general-risks", "payload": {"risk_register": {"top_risks": [{
                    "trigger": "opex_above_60_percent_of_revenue", "severity": "high",
                    "mitigation": "Reduce fixed OPEX or increase validated revenue capacity.",
                }]}}},
                {"section_id": "09-timeline", "payload": {"milestones": [{
                    "phase_id": "setup", "owner_role": "Project Manager",
                    "estimated_duration_days": 17, "exit_criteria": ["project_scope_signed"],
                }]}},
            ],
        }
        original = deepcopy(projection)
        with tempfile.TemporaryDirectory() as directory:
            for locale, expected in (
                ("ar", ["الإيراد الشهري", "المصروفات التشغيلية مرتفعة", "اعتماد نطاق المشروع"]),
                ("en", ["Monthly revenue", "Operating costs are high", "Approve project scope"]),
            ):
                with self.subTest(locale=locale):
                    outputs = [render_funder_report_html(projection, locale)]
                    deck = export_funder_report_pptx(projection, Path(directory) / f"report-{locale}.pptx", locale)
                    with zipfile.ZipFile(deck) as archive:
                        slide_text = []
                        for name in archive.namelist():
                            if name.startswith("ppt/slides/slide") and name.endswith(".xml"):
                                root = ElementTree.fromstring(archive.read(name))
                                slide_text.extend(node.text or "" for node in root.iter("{http://schemas.openxmlformats.org/drawingml/2006/main}t"))
                        outputs.append(" ".join(slide_text))
                    try:
                        from docx import Document
                    except ModuleNotFoundError:
                        Document = None
                    if Document is not None:
                        word = export_funder_report_docx(projection, Path(directory) / f"report-{locale}.docx", locale)
                        document = Document(word)
                        outputs.append(" ".join(
                            [paragraph.text for paragraph in document.paragraphs] +
                            [cell.text for table in document.tables for row in table.rows for cell in row.cells]
                        ))
                    for output in outputs:
                        self.assertIn("12345.6789", output)
                        for phrase in expected:
                            self.assertIn(phrase, output)
                        for private in ("PRIVATE-SNAPSHOT", "PRIVATE-HASH", "monthly_revenue", "project_scope_signed", "DRAFT_INTERNAL"):
                            self.assertNotIn(private, output)
        self.assertEqual(original, projection)

    def test_customer_report_groups_reject_non_numeric_metric_payloads(self) -> None:
        projection = {"sections": [{"section_id": "02-executive-summary", "payload": {
            "kpis": [{"output_id": "monthly_revenue", "value": "PRIVATE-runtime", "unit": "PRIVATE_UNIT", "status": "not_ready"}],
        }}]}
        groups = customer_report_groups(projection, "en")
        self.assertEqual("—", groups[1]["rows"][0][1])
        self.assertEqual("—", groups[1]["rows"][0][2])
        self.assertNotIn("PRIVATE", str(groups))

    def test_reference_profiles_return_explainable_missing_requirements(self) -> None:
        repo = self.make_repo()
        project = repo.create_project({"name": "ملف جهة", "inputs": {}})
        overview, report = api.build_overview(project, repo)

        self.assertIn("BANK-SME-BASE-V1", profile_ids())
        result = evaluate_funding_readiness(report["funder_report"], "BANK-SME-BASE-V1")
        self.assertEqual(result["snapshot_id"], overview["snapshot"]["snapshot_id"])
        self.assertEqual(result["status"], "DRAFT_INTERNAL")
        self.assertIn("financial_projection", result["missing_requirements"])
        self.assertIn("لا يمثل قبولاً", result["acceptance_disclaimer"])

    def test_sector_profiles_are_scoped_and_reference_only(self) -> None:
        profiles = sector_profile_catalog()
        self.assertGreaterEqual(len(profiles), 4)
        retail = next(profile for profile in profiles if profile["profile_id"] == "SECTOR-RETAIL-V1")
        self.assertEqual(retail["profile_status"], "reference_only")
        self.assertEqual(retail["reviewed_at"], "2026-07-20")
        self.assertTrue(retail["not_covered_ar"])
        self.assertIn("lender_policy", retail["not_locally_verifiable"])

    def test_release_requires_matching_human_review_and_profile_readiness(self) -> None:
        repo = self.make_repo()
        project = repo.create_project({"name": "إصدار محلي", "inputs": {}})
        _overview, report = api.build_overview(project, repo)
        projection = report["funder_report"]

        blocked = build_release_record(projection, None)
        self.assertEqual(blocked["release_state"], "REVIEW_REQUIRED")
        self.assertIn("human_review_missing", blocked["blocking_reasons"])

        review = {"review_id": "review-1", "snapshot_id": projection["snapshot_id"], "reviewer": "reviewer", "decision": "approved_local", "created_at": "2026-07-20T00:00:00+00:00"}
        released = build_release_record(projection, review)
        self.assertEqual(released["release_state"], "REVIEW_REQUIRED")
        validate_release_record(projection, released)


if __name__ == "__main__":
    unittest.main()
