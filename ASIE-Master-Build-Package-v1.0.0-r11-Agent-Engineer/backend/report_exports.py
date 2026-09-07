"""File exporters for a persisted funder-report projection.

Exporters consume the already-built projection. They do not calculate finance,
change Snapshot state, or infer missing accounting values.
"""

from __future__ import annotations

from pathlib import Path
import os
import shutil
import subprocess
import tempfile
from zipfile import ZIP_DEFLATED, ZipFile
from typing import Any

from backend.customer_presentation import business_text, normalize_locale, section_title, status_text, text
from backend.customer_report_content import customer_report_groups

try:
    from docx.shared import RGBColor
except ModuleNotFoundError:  # The bundled document runtime is loaded only for export.
    RGBColor = Any  # type: ignore[misc,assignment]


NAVY = "172554"
BLUE = "2563EB"
MUTED = "64748B"


def export_funder_report_pptx(projection: dict[str, Any], output_path: str | Path, locale: str = "ar") -> Path:
    """Create a dependency-free localized PPTX from the persisted projection."""
    locale = normalize_locale(locale)
    from xml.sax.saxutils import escape
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    sections = projection.get("sections", [])
    # Paginate customer content rather than exporting only section names.
    import textwrap
    slides = [(text("report_title", locale), status_text(projection.get("readiness_status"), locale))]

    def add_slides(title: str, lines: list[str]) -> None:
        wrapped = [part for line in lines for part in (textwrap.wrap(str(line), width=78) or [""])]
        for start in range(0, max(1, len(wrapped)), 8):
            slides.append((title, "\n".join(wrapped[start:start + 8])))

    add_slides(text("readiness", locale), [text("notice", locale)])
    for group in customer_report_groups(projection, locale):
        lines = [
            " · ".join(f"{header}: {value}" for header, value in zip(group["headers"], row))
            for row in group["rows"]
        ]
        add_slides(group["title"], lines or [group["empty"]])
    add_slides(text("study_structure", locale), [section_title(row, locale) for row in sections])
    add_slides(text("missing_items", locale), [business_text(gap, locale) for gap in projection.get("gaps", [])] or [text("none", locale)])
    add_slides(text("report_status", locale), [text("saved", locale)])
    def text_shape(name: str, content: str, y: int, size: int) -> str:
        body_direction = ' rtlCol="1"' if locale == "ar" else ""
        paragraph_direction = ' rtl="1" algn="r"' if locale == "ar" else ' algn="l"'
        language = "ar-SA" if locale == "ar" else "en-US"
        paragraphs = "".join(
            f'<a:p><a:pPr{paragraph_direction}/><a:r><a:rPr lang="{language}" sz="{size * 100}"/>'
            f'<a:t>{escape(line)}</a:t></a:r></a:p>'
            for line in (str(content).splitlines() or [""])
        )
        height = 900000 if name == "Title" else 4000000
        return f'<p:sp><p:nvSpPr><p:cNvPr id="{y}" name="{name}"/><p:cNvSpPr/><p:nvPr/></p:nvSpPr><p:spPr><a:xfrm><a:off x="900000" y="{y}"/><a:ext cx="10300000" cy="{height}"/></a:xfrm><a:prstGeom prst="rect"><a:avLst/></a:prstGeom></p:spPr><p:txBody><a:bodyPr{body_direction}/><a:lstStyle/>{paragraphs}</p:txBody></p:sp>'
    def slide_xml(title: str, body: str) -> str:
        return '<?xml version="1.0" encoding="UTF-8" standalone="yes"?><p:sld xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main"><p:cSld><p:spTree><p:nvGrpSpPr><p:cNvPr id="1" name=""/><p:cNvGrpSpPr/><p:nvPr/></p:nvGrpSpPr><p:grpSpPr/>' + text_shape("Title", title, 700000, 28) + text_shape("Body", body, 2500000, 16) + '</p:spTree></p:cSld><p:clrMapOvr><a:masterClrMapping/></p:clrMapOvr></p:sld>'
    content_types = '<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"><Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/><Default Extension="xml" ContentType="application/xml"/><Override PartName="/ppt/presentation.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.presentation.main+xml"/><Override PartName="/ppt/slideMasters/slideMaster1.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.slideMaster+xml"/><Override PartName="/ppt/slideLayouts/slideLayout1.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.slideLayout+xml"/>' + ''.join(f'<Override PartName="/ppt/slides/slide{i}.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.slide+xml"/>' for i in range(1, len(slides) + 1)) + '</Types>'
    rels = '<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
    with ZipFile(output, "w", compression=ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", content_types)
        archive.writestr("_rels/.rels", rels + '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="ppt/presentation.xml"/></Relationships>')
        archive.writestr("ppt/presentation.xml", '<?xml version="1.0" encoding="UTF-8" standalone="yes"?><p:presentation xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main"><p:sldMasterIdLst><p:sldMasterId id="2147483648" r:id="rId1"/></p:sldMasterIdLst><p:sldIdLst>' + ''.join(f'<p:sldId id="{255 + i}" r:id="rId{i + 1}"/>' for i in range(1, len(slides) + 1)) + '</p:sldIdLst><p:sldSz cx="12192000" cy="6858000" type="screen16x9"/></p:presentation>')
        archive.writestr("ppt/_rels/presentation.xml.rels", rels + '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slideMaster" Target="slideMasters/slideMaster1.xml"/>' + ''.join(f'<Relationship Id="rId{i + 1}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slide" Target="slides/slide{i}.xml"/>' for i in range(1, len(slides) + 1)) + '</Relationships>')
        archive.writestr("ppt/slideMasters/slideMaster1.xml", '<?xml version="1.0" encoding="UTF-8" standalone="yes"?><p:sldMaster xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main"><p:cSld name="Master"><p:spTree><p:nvGrpSpPr><p:cNvPr id="1" name=""/><p:cNvGrpSpPr/><p:nvPr/></p:nvGrpSpPr><p:grpSpPr/></p:spTree></p:cSld><p:sldLayoutIdLst><p:sldLayoutId id="1" r:id="rId1"/></p:sldLayoutIdLst></p:sldMaster>')
        archive.writestr("ppt/slideMasters/_rels/slideMaster1.xml.rels", rels + '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slideLayout" Target="../slideLayouts/slideLayout1.xml"/></Relationships>')
        archive.writestr("ppt/slideLayouts/slideLayout1.xml", '<?xml version="1.0" encoding="UTF-8" standalone="yes"?><p:sldLayout xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main"><p:cSld name="Blank"><p:spTree><p:nvGrpSpPr><p:cNvPr id="1" name=""/><p:cNvGrpSpPr/><p:nvPr/></p:nvGrpSpPr><p:grpSpPr/></p:spTree></p:cSld></p:sldLayout>')
        for index, (title, body) in enumerate(slides, 1):
            archive.writestr(f"ppt/slides/slide{index}.xml", slide_xml(title, body))
            archive.writestr(f"ppt/slides/_rels/slide{index}.xml.rels", rels + '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slideLayout" Target="../slideLayouts/slideLayout1.xml"/></Relationships>')
    return output


def _set_paragraph_direction(paragraph: Any, locale: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    locale = normalize_locale(locale)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if locale == "ar" else WD_ALIGN_PARAGRAPH.LEFT
    ppr = paragraph._p.get_or_add_pPr()
    bidi = ppr.find(qn("w:bidi"))
    if locale == "ar" and bidi is None:
        bidi = OxmlElement("w:bidi")
        ppr.append(bidi)
    elif locale == "en" and bidi is not None:
        ppr.remove(bidi)


def _set_font(run: Any, *, size: float = 11, color: Any = "1E293B", bold: bool = False) -> None:
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color if isinstance(color, str) else str(color))
    run.bold = bold


def _shade(cell: Any, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_text(cell: Any, text: Any, *, locale: str, bold: bool = False, color: Any = "1E293B") -> None:
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    cell.text = ""
    paragraph = cell.paragraphs[0]
    _set_paragraph_direction(paragraph, locale)
    run = paragraph.add_run(str(text if text not in (None, "") else "—"))
    _set_font(run, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _table(doc: Any, headers: list[str], rows: list[list[Any]], *, locale: str) -> Any:
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.shared import Inches
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT if normalize_locale(locale) == "ar" else WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for index, header in enumerate(headers):
        _set_cell_text(table.rows[0].cells[index], header, locale=locale, bold=True, color=RGBColor(255, 255, 255))
        _shade(table.rows[0].cells[index], "172554")
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            _set_cell_text(cells[index], value, locale=locale)
            if len(table.rows) % 2 == 0:
                _shade(cells[index], "F8FAFC")
    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(6.5 / len(headers))
    return table


def export_funder_report_docx(projection: dict[str, Any], output_path: str | Path, locale: str = "ar") -> Path:
    """Create a read-only localized DOCX from a funder projection."""
    locale = normalize_locale(locale)
    try:
        from docx import Document
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor
    except ModuleNotFoundError as exc:
        raise RuntimeError("DOCX export requires the bundled document runtime") from exc
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    normal.font.size = Pt(10.5)

    header = section.header.paragraphs[0]
    _set_paragraph_direction(header, locale)
    run = header.add_run(f"ASIE | {text('report_title', locale)}")
    _set_font(run, size=9, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    _set_paragraph_direction(footer, locale)
    run = footer.add_run(f"{text('report_status', locale)}: {text('saved', locale)}")
    _set_font(run, size=8, color=MUTED)

    title = doc.add_paragraph()
    _set_paragraph_direction(title, locale)
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run(text("report_title", locale))
    _set_font(run, size=24, color=NAVY, bold=True)
    subtitle = doc.add_paragraph()
    _set_paragraph_direction(subtitle, locale)
    run = subtitle.add_run(f"{text('report_status', locale)}: {status_text(projection.get('readiness_status'), locale)}")
    _set_font(run, size=11, color=BLUE, bold=True)

    note = doc.add_paragraph()
    _set_paragraph_direction(note, locale)
    run = note.add_run(text("notice", locale))
    _set_font(run, size=10, color="7C2D12", bold=True)

    for group in customer_report_groups(projection, locale):
        heading = doc.add_heading(group["title"], level=1)
        _set_paragraph_direction(heading, locale)
        if group["rows"]:
            _table(doc, group["headers"], group["rows"], locale=locale)
        else:
            paragraph = doc.add_paragraph(group["empty"])
            _set_paragraph_direction(paragraph, locale)

    h = doc.add_heading(text("report_status", locale), level=1)
    _set_paragraph_direction(h, locale)
    _set_font(h.runs[0], size=16, color=BLUE, bold=True)
    summary = next((row for row in projection.get("sections", []) if row.get("section_id") == "02-executive-summary"), {})
    payload = summary.get("payload") or {}
    _table(doc, [text("requirement", locale), text("status", locale)], [[text("readiness", locale), status_text(projection.get("readiness_status"), locale)], [text("report_status", locale), text("saved", locale)]], locale=locale)

    h = doc.add_heading(text("readiness", locale), level=1)
    _set_paragraph_direction(h, locale)
    profile = projection.get("profile_readiness") or {}
    _table(doc, [text("requirement", locale), text("status", locale), text("reason", locale)], [[business_text(row.get("label"), locale), status_text(row.get("status"), locale), business_text(row.get("reason"), locale)] for row in profile.get("checks", [])] or [["—", status_text("not_ready", locale), text("no_checks", locale)]], locale=locale)

    h = doc.add_heading(text("study_structure", locale), level=1)
    _set_paragraph_direction(h, locale)
    rows = [[section_title(row, locale), status_text(row.get("status"), locale)] for row in projection.get("sections", [])]
    _table(doc, [text("requirement", locale), text("status", locale)], rows, locale=locale)

    h = doc.add_heading(text("financial_outlook", locale), level=1)
    _set_paragraph_direction(h, locale)
    financial = next((row for row in projection.get("sections", []) if row.get("section_id") == "14-financial-expectations"), {})
    statements = (financial.get("payload") or {}).get("statements") or {}
    years = ((statements.get("income_statement") or {}).get("years") or [])
    _table(doc, [text("year", locale), text("revenue", locale), text("gross_profit", locale), text("ebitda", locale), text("ebit", locale), text("operating_cashflow", locale)], [[row.get("year"), row.get("revenue"), row.get("gross_profit"), row.get("ebitda"), row.get("ebit"), row.get("net_operating_cashflow")] for row in years] or [["—", text("no_financials", locale), "—", "—", "—", "—"]], locale=locale)

    h = doc.add_heading(text("missing_items", locale), level=1)
    _set_paragraph_direction(h, locale)
    for gap in projection.get("gaps", []):
        paragraph = doc.add_paragraph(style="List Bullet")
        _set_paragraph_direction(paragraph, locale)
        run = paragraph.add_run(business_text(gap, locale))
        _set_font(run, size=10)

    h = doc.add_heading(text("report_status", locale), level=1)
    _set_paragraph_direction(h, locale)
    evidence = projection.get("evidence") or {}
    _table(doc, [text("requirement", locale), text("status", locale)], [[text("readiness", locale), status_text(projection.get("readiness_status"), locale)], [text("report_status", locale), text("saved", locale)]], locale=locale)
    doc.save(output)
    return output


def export_funder_report_pdf(
    projection: dict[str, Any], output_path: str | Path, renderer_path: str | Path | None = None, locale: str = "ar"
) -> Path:
    """Print the canonical Arabic HTML projection with a server-side renderer.

    The client browser is never involved. ``renderer_path`` (or the
    ``ASIE_PDF_RENDERER`` setting) identifies the pinned headless renderer in
    production. Chrome/Edge discovery is retained only as a local development
    fallback; Firefox/Safari/other client browsers can download the result
    identically through the API.
    """
    from backend.funder_report import render_funder_report_html

    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    configured = renderer_path or os.environ.get("ASIE_PDF_RENDERER")
    browser = str(configured) if configured else (shutil.which("chrome") or shutil.which("msedge"))
    if browser is None:
        candidates = (
            Path(r"C:\Program Files\Google\Chrome\Application\chrome.exe"),
            Path(r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe"),
        )
        browser = next((str(path) for path in candidates if path.exists()), None)
    if browser is None:
        raise RuntimeError("PDF export requires a configured server-side PDF renderer")
    with tempfile.TemporaryDirectory(prefix="asie-funder-pdf-") as temp_dir:
        html_path = Path(temp_dir) / "funder-report.html"
        html_path.write_text(render_funder_report_html(projection, locale=locale), encoding="utf-8")
        command = [
            browser,
            "--headless",
            "--disable-gpu",
            "--no-sandbox",
            "--no-pdf-header-footer",
            f"--print-to-pdf={output.resolve()}",
            html_path.resolve().as_uri(),
        ]
        subprocess.run(command, check=True, capture_output=True, text=True)
    if not output.exists() or output.stat().st_size == 0:
        raise RuntimeError("PDF renderer returned without creating a file")
    return output
